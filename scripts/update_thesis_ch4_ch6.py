from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
THESIS_DIR = ROOT / "论文"
FIG_DIR = ROOT / "docs" / "_thesis_generated"
SOURCE_NAME = "10论文_基于LLM与Flask的股票量化分析系统 .docx"
OUTPUT_NAME = "10论文_基于LLM与Flask的股票量化分析系统_第4章第6章修订版_20260429.docx"


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.text = text


def add_centered_picture(anchor: Paragraph, image_path: Path, caption: str) -> Paragraph:
    image_para = insert_paragraph_after(anchor)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(image_path), width=Inches(5.8))

    caption_para = insert_paragraph_after(image_para, caption, style="Caption")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return caption_para


def add_table_after(doc: Document, anchor: Paragraph, caption: str, rows: list[list[str]]) -> Paragraph:
    caption_para = insert_paragraph_after(anchor, caption, style="图表标题")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
            for p in table.cell(i, j).paragraphs:
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para._p.addnext(table._tbl)

    spacer = insert_paragraph_after(caption_para)
    table._tbl.addnext(spacer._p)
    return spacer


def update_chapter_four(doc: Document) -> None:
    p41 = find_paragraph(doc, "4.1系统总体架构设计")
    p42 = find_paragraph(doc, "4.2 系统功能模块设计")
    p421 = find_paragraph(doc, "4.2.1 大盘速览")
    p422 = find_paragraph(doc, "4.2.2 股票列表")
    p423 = find_paragraph(doc, "4.2.3 个股分析")
    p424 = find_paragraph(doc, "4.2.4 股票筛选")
    p425 = find_paragraph(doc, "4.2.5 实时分析")
    p426 = find_paragraph(doc, "4.2.6 AI助手")
    p427 = find_paragraph(doc, "4.2.7 个人中心")
    p428 = find_paragraph(doc, "4.2.8 管理员后台")
    p43 = find_paragraph(doc, "4.3 数据库设计")
    p441 = find_paragraph(doc, "4.4.1 实体图")
    p442 = find_paragraph(doc, "4.4.2 ER图")
    p443 = find_paragraph(doc, "4.4.3 数据库表设计")

    anchor = p41
    anchor = insert_paragraph_after(
        anchor,
        "本系统采用典型的 B/S 架构，以浏览器作为统一访问入口，前端使用 Jinja2 模板与原生 JavaScript 完成页面渲染和交互，后端基于 Flask 提供页面路由、JSON 接口、权限控制和异常处理能力。用户在首页、股票列表、个股分析、实时监控和 AI 助手等页面发起请求后，系统会根据不同场景分别进入页面渲染链路或 API 数据链路，从而兼顾交互体验与开发维护成本。",
    )
    anchor = insert_paragraph_after(
        anchor,
        "在服务层设计上，系统按照“页面层、接口层、业务服务层、数据模型层、外部数据接入层”进行分层。其中，app/main 负责页面入口，app/api 负责股票、分析、新闻、监控和 AI 等接口，app/routes 负责认证与后台管理，app/services 封装市场概览、实时监控、新闻聚合、大模型调用、用户行为记录等核心业务逻辑，app/models 负责 ORM 模型与数据库映射。该分层方式能够降低模块耦合度，使系统在功能扩展时更容易定位责任边界。",
    )
    anchor = insert_paragraph_after(
        anchor,
        "在基础设施方面，MySQL 作为主业务数据库，Redis 作为高频接口缓存并支持降级到内存缓存，Tushare 与 AkShare 共同承担金融数据接入，LLMService 通过 OpenAI 兼容接口统一封装大模型能力。对于实时行情、缓存、外部接口等不稳定依赖，系统均提供回退或降级方案，保证股票查询、监控看板和历史分析等核心链路在部分依赖异常时仍可继续运行。",
    )
    fig1 = FIG_DIR / "fig4_1_architecture.png"
    if fig1.exists():
        anchor = add_centered_picture(anchor, fig1, "图4-1 系统总体架构图")

    anchor = p42
    anchor = insert_paragraph_after(
        anchor,
        "结合股票投资新手的使用需求与管理员的维护需求，系统功能模块被划分为用户端分析模块与后台管理模块两大部分。用户端围绕“查行情、做分析、看回测、盯盘面、问 AI、管记录”展开，后台端围绕“管用户、看日志、查数据、做自检”展开，各模块之间通过统一会话机制和数据模型实现联动。",
    )
    fig2 = FIG_DIR / "fig4_2_modules.png"
    if fig2.exists():
        anchor = add_centered_picture(anchor, fig2, "图4-2 系统功能模块图")

    insert_paragraph_after(
        p421,
        "大盘速览模块主要服务于首页看板，负责整合上证指数、深证成指、行业热度、成交额排行等概览信息。该模块优先从实时数据源获取市场快照，在非交易时段或接口异常时回退到最近交易日缓存结果，为用户提供一个低门槛、可快速理解市场状态的入口。",
    )
    insert_paragraph_after(
        p422,
        "股票列表模块面向全市场股票信息查询场景，支持按股票代码、名称、行业和地域等条件进行筛选，并提供分页浏览能力。用户从该模块可以进一步进入股票详情页，形成“列表检索 - 详情分析 - 加入自选”的连续研究路径。",
    )
    insert_paragraph_after(
        p423,
        "个股分析模块负责展示单只股票的核心研究信息，包括历史行情、K 线走势、均线与 MACD 等技术指标、资金流向、筹码分布与基础属性信息。该模块是系统中数据最集中的业务节点，也是连接股票列表、策略回测与个人中心记录沉淀的重要桥梁。",
    )
    insert_paragraph_after(
        p424,
        "股票筛选模块采用多条件组合过滤方式，既支持行业、地域、市值、估值、换手率、量比等固定条件，也支持 MA5>MA10、PE<PB 等动态字段比较条件。筛选结果会返回结构化股票列表，方便用户将定量条件快速转化为候选研究池。",
    )
    insert_paragraph_after(
        p425,
        "实时分析模块面向盘中观察场景，集成市场总览、自选股监控、涨跌排行、分时走势和异常波动等功能。系统通过 RealtimeMonitorService 统一管理行情获取、信号生成和降级逻辑，当分钟级数据不可用时可自动回退到最近交易日数据，以提升监控页面的稳定性。",
    )
    insert_paragraph_after(
        p426,
        "AI 助手模块为用户提供自然语言问答能力，支持围绕个股、技术指标、交易策略和风险控制等主题进行对话。页面层允许用户直接进入助手界面查看模型配置与状态，而正式对话、历史会话读取和会话管理则需要在登录状态下通过 AIConversationService 与 LLMService 协同完成。",
    )
    insert_paragraph_after(
        p427,
        "个人中心模块用于统一管理用户账户资料、自选股、分析记录和 AI 问答记录。该模块不仅承担基本的账户维护职能，还通过统计信息和最近记录展示帮助用户保留研究轨迹，提高后续复盘与持续学习的便利性。",
    )
    insert_paragraph_after(
        p428,
        "管理员后台模块面向系统运维与治理场景，包含后台总览、用户管理、日志审计、数据中心与系统自检等子功能。管理员可以在浏览器端完成用户状态调整、角色分配、分钟级数据同步和依赖组件连通性检查，从而降低系统维护的操作门槛。",
    )

    anchor = p43
    anchor = insert_paragraph_after(
        anchor,
        "数据库设计以 MySQL 为核心，遵循“业务对象清晰、时序数据分离、用户行为可追踪、系统运维可审计”的设计原则。考虑到系统既要承载用户侧交互，又要存储大规模股票行情与技术指标，因此在逻辑上将数据划分为用户与行为数据、AI 会话数据、股票基础与行情数据、技术分析衍生数据以及系统日志数据等多个子域。",
    )
    anchor = insert_paragraph_after(
        anchor,
        "对于高频查询的股票数据，系统重点围绕 ts_code、trade_date 等字段组织索引和联合查询路径，以支撑股票详情、条件选股、策略回测和实时监控的读性能需求；对于用户行为和后台审计数据，则重点围绕 user_id、created_at、action_type 等字段实现按用户、按时间和按操作类型的检索能力。这种设计方式既满足了分析类场景的数据读取效率，也为后续扩展用户资产管理和风险提示功能保留了空间。",
    )

    heading_44 = p441.insert_paragraph_before("4.4 核心实体与数据表设计", style="Heading 2")
    insert_paragraph_after(
        heading_44,
        "考虑到论文篇幅，本节重点选取用户账户、系统日志、自选股、历史问答、股票基础信息、日线行情、技术指标、交易信号与模拟持仓等能够支撑当前主业务链路的核心实体和数据表进行说明。",
    )

    insert_paragraph_after(
        p441,
        "系统的核心实体既包括用户与管理员等主体对象，也包括股票、行情、技术指标、问答记录和系统日志等业务对象。各实体共同支撑“用户研究 - 数据分析 - 历史留痕 - 后台维护”的完整业务闭环。",
    )

    entity_updates = {
        "（1）用户实体图": "（1）用户实体：用户表保存用户名、邮箱、密码哈希、账户角色、账户状态、最近登录时间等字段，是登录认证、权限校验和个人中心展示的基础。",
        "（2）系统操作日志实体图": "（2）系统操作日志实体：系统日志表记录操作用户、操作类型、说明信息、来源 IP、执行状态和创建时间，用于后台审计与问题排查。",
        "（3）用户自选股实体图": "（3）用户自选股实体：自选股表描述用户与股票之间的关注关系，包含股票代码、股票名称、来源、备注和排序等信息，用于个人研究池管理。",
        "（4）智能问答记录实体图": "（4）智能问答记录实体：历史问答实体保存用户提问内容、AI 回复内容、摘要信息和创建时间，是 AI 助手历史沉淀与个人中心展示的重要依据。",
        "（5）股票基础信息实体图": "（5）股票基础信息实体：股票基础信息表保存股票代码、简称、上市日期、所属行业、地域和市场属性等静态信息，为列表检索和详情展示提供基础数据。",
        "（6）日线行情实体图": "（6）日线行情实体：日线行情表按股票代码和交易日保存开盘价、收盘价、最高价、最低价、成交量和成交额等时序数据，是图表展示和回测计算的核心输入。",
        "（7）每日技术与基本面指标实体图": "（7）每日技术与基本面指标实体：该实体保存换手率、市盈率、市净率、股息率等估值与交易指标，为选股和分析页面提供结构化基础。",
        "（8）股票技术因子数据": "（8）股票技术因子实体：技术因子表保存 MA、MACD、RSI、KDJ、BOLL、CCI 等派生指标，供个股分析、实时信号生成和策略回测复用。",
        "（9）交易信号与回测记录实体图": "（9）交易信号与回测记录实体：交易信号表记录策略触发时刻、信号类型、价格和说明信息；回测结果则保存策略名称、时间区间、收益率、回撤和交易明细等统计结果。",
        "（10）模拟持仓管理实体图": "（10）模拟持仓管理实体：模拟持仓表用于记录组合编号、股票代码、持仓数量、成本价、当前价格和盈亏情况，为后续扩展虚拟组合管理能力提供数据基础。",
    }
    for prefix, text in entity_updates.items():
        set_paragraph_text(find_paragraph(doc, prefix), text)

    insert_paragraph_after(
        p442,
        "从实体关系上看，用户与自选股、分析记录、聊天记录和回测结果之间均为一对多关系；股票基础信息与日线行情、每日指标、技术因子、分钟线数据和异动数据之间也形成一对多关系；系统日志通过用户标识与后台或普通用户行为建立关联。这样的 ER 关系能够较好地覆盖当前系统“账号 - 研究行为 - 股票数据 - 维护审计”的主线业务。",
    )

    insert_paragraph_after(
        p443,
        "在数据库表设计层面，系统重点展示 10 张与当前主业务链路最紧密的数据表。各数据表既要满足页面展示与接口查询的需要，也要支持后续的数据同步、统计分析和后台审计等扩展场景。",
    )

    table_label_updates = {
        "（1）用户表设计如表 4-1": "（1）用户账户表设计如表 4-1 所示。",
        "表4-1 用户表": "表4-1 用户表 (user_account)",
        "（2）系统操作日志表设计如表 4-2": "（2）系统操作日志表设计如表 4-2 所示。",
        "表4-2 系统操作日志表": "表4-2 系统操作日志表 (system_log)",
        "（3）用户自选股表设计如表 4-3": "（3）用户自选股表设计如表 4-3 所示。",
        "表4-3 用户自选股表": "表4-3 用户自选股表 (user_watchlist)",
        "（4）智能问答记录表设计如表 4-4": "（4）智能问答记录表设计如表 4-4 所示。",
        "表4-4 智能问答记录表": "表4-4 智能问答记录表 (user_chat_history)",
        "（5）股票基础信息表设计如表 4-5": "（5）股票基础信息表设计如表 4-5 所示。",
        "表4-5 股票基础信息表": "表4-5 股票基础信息表 (stock_basic)",
        "（6）日线行情表设计如表 4-6": "（6）日线行情表设计如表 4-6 所示。",
        "表4-6 日线行情表": "表4-6 日线行情表 (stock_daily_history)",
        "（7）每日技术与基本面指标表设计如表 4-7": "（7）每日技术与基本面指标表设计如表 4-7 所示。",
        "表4-7 每日技术与基本面指标表": "表4-7 每日技术与基本面指标表 (stock_daily_basic)",
        "（8）股票技术因子数据表设计如表4-8": "（8）股票技术因子数据表设计如表 4-8 所示。",
        "表4-8 股票技术因子数据表": "表4-8 股票技术因子数据表 (stock_factor)",
        "（10）交易信号与回测记录表设计如表4-10": "（9）交易信号表设计如表 4-9 所示。",
        "表4- 9 交易信号与回测记录数据表": "表4-9 交易信号表 (trading_signals)",
        "（11）模拟持仓管理表设计如表4-11": "（10）模拟持仓管理表设计如表 4-10 所示。",
        "表4-10 模拟持仓比数据表": "表4-10 模拟持仓管理表 (portfolio_positions)",
    }
    for prefix, text in table_label_updates.items():
        set_paragraph_text(find_paragraph(doc, prefix), text)


def update_chapter_six(doc: Document) -> None:
    p61 = find_paragraph(doc, "6.1 测试环境与测试方法")
    p62 = find_paragraph(doc, "6.2 主要功能测试")
    p621 = find_paragraph(doc, "（1）访问控制测试")
    p622 = find_paragraph(doc, "（2）核心数据接口测试")
    p623 = find_paragraph(doc, "（3）个人功能测试")
    p624 = find_paragraph(doc, "（4）管理维护功能测试")
    p63 = find_paragraph(doc, "6.3 测试结果分析")

    p61_texts = [
        "为验证系统主要功能的可用性与稳定性，本文于 2026 年 4 月 29 日在本地开发环境中，对页面访问链路、权限控制链路、核心数据接口和后台自检能力进行了集中验证。测试方式包括浏览器页面访问检查与 Flask `test_client` 接口冒烟测试两类，前者主要观察页面跳转与交互路径，后者主要核验返回状态码、JSON 结构以及降级处理是否符合预期。",
        "考虑到本系统同时依赖 MySQL、Redis、Tushare、AkShare 和大模型服务等外部组件，测试过程不仅关注页面是否能正常打开，也关注在部分组件不可用时系统是否能够自动回退到可接受的运行模式。因此，本章重点选择了股票列表、实时监控、AI 助手、个人中心、后台总览、数据中心和系统自检等最能代表主业务链路的场景进行验证。",
    ]
    # Replace the first paragraph body beneath 6.1 and insert one more.
    p_env_1 = find_paragraph(doc, "为验证系统的可用性与稳定性")
    set_paragraph_text(p_env_1, p61_texts[0])
    p_env_2 = find_paragraph(doc, "测试方法主要分为两类")
    set_paragraph_text(p_env_2, p61_texts[1])
    p_env_3 = find_paragraph(doc, "测试过程中，重点关注")
    set_paragraph_text(
        p_env_3,
        "本次测试重点覆盖匿名访问控制、普通用户功能可用性、管理员后台权限、股票与监控接口返回结果，以及系统在缓存或外部行情组件异常时的容错表现。",
    )

    env_rows = [
        ["项目", "配置", "说明"],
        ["测试时间", "2026年4月29日", "本地集中验证时间"],
        ["操作系统", "Windows 本地开发环境", "论文当前演示环境"],
        ["Python", "3.11.1", "应用运行时"],
        ["Web 框架", "Flask 1.1.4", "系统后端主框架"],
        ["数据库", "MySQL 8.1.0", "业务主库"],
        ["缓存", "Redis + 内存降级", "Redis 未启动时自动回退"],
        ["主要数据源", "Tushare Pro、AkShare / 新浪快照", "股票历史与实时数据来源"],
        ["测试方式", "页面访问 + Flask test_client 冒烟测试", "覆盖页面和接口两类场景"],
    ]
    anchor = add_table_after(doc, p_env_3, "表6-1 测试环境配置", env_rows)

    insert_paragraph_after(
        p62,
        "结合页面验证结果和接口返回结果，本次测试主要从访问控制、用户功能、核心接口与后台运维四个角度进行归纳，测试结果如表 6-2 和表 6-3 所示。",
    )

    p621_body = find_paragraph(doc, "对首页、股票列表、AI助手、实时监控等页面进行未登录访问测试时")
    set_paragraph_text(
        p621_body,
        "在访问控制方面，匿名用户访问首页可以直接进入系统入口页，访问股票列表和实时监控页面时会被重定向至登录页；AI 助手页面本身允许直接打开，但正式发起会话和读取历史记录仍要求用户登录。普通用户登录后可以正常访问个人中心、策略回测和实时监控等页面；管理员登录后能够正常访问后台总览、用户管理、日志管理、数据中心和系统自检页面，说明系统已实现前后台分级授权。",
    )
    p622_body = find_paragraph(doc, "对健康检查接口、股票列表接口、监控面板接口和实时排名接口进行调用测试时")
    set_paragraph_text(
        p622_body,
        "在核心数据接口方面，系统健康检查接口、股票列表接口、监控面板接口、实时排名接口和条件选股接口均成功返回状态码 200，并能够输出健康状态、股票列表、监控面板、实时排名以及筛选结果等结构化数据。其中，空条件调用选股接口时返回了 5388 条股票记录，说明筛选链路、数据库查询与返回封装能够正常工作。",
    )
    p623_body = find_paragraph(doc, "用户登录后访问个人中心页面时")
    set_paragraph_text(
        p623_body,
        "在个人功能方面，普通用户登录后访问个人中心页面时，系统能够正确展示自选股数量、分析记录数量和 AI 问答数量，并可读取最近的自选股、分析记录和问答记录。同时，AI 历史会话接口、分析记录接口和回测历史接口均正常返回历史数据，说明用户研究痕迹的存储与读取链路已经具备可用性。",
    )
    p624_body = find_paragraph(doc, "管理员调用系统自检接口后")
    set_paragraph_text(
        p624_body,
        "在后台维护方面，管理员访问后台各页面均返回 200 状态码。进一步调用系统自检接口后，MySQL 与 AkShare 组件状态为正常，Redis 因本地服务未启动被标记为 disabled，系统随后自动回退到内存缓存；Tushare 在本次环境中因本地令牌文件权限问题被标记为 error。虽然整体状态显示为 degraded，但股票列表、监控排行和后台页面仍可继续工作，说明系统具备一定的降级和容错能力。",
    )

    case_rows = [
        ["编号", "测试场景", "预期结果", "实际结果"],
        ["T1", "匿名访问系统首页", "页面可正常打开", "通过，HTTP 200"],
        ["T2", "匿名访问股票列表与实时监控页面", "跳转到登录页", "通过，HTTP 302"],
        ["T3", "匿名访问 AI 助手页面", "页面可打开，正式会话需登录", "通过，页面 HTTP 200"],
        ["T4", "普通用户访问个人中心、策略回测、实时监控", "页面正常展示", "通过，均为 HTTP 200"],
        ["T5", "管理员访问后台总览、用户管理、日志、数据中心、自检页", "页面正常展示", "通过，均为 HTTP 200"],
        ["T6", "调用健康检查、股票列表、监控面板、实时排名、条件选股接口", "返回 200 且数据结构完整", "通过，均返回结构化 JSON"],
        ["T7", "调用 AI 历史会话、分析记录、回测历史接口", "可读取历史记录", "通过，均返回用户历史数据"],
    ]
    anchor = add_table_after(doc, p624_body, "表6-2 主要功能与接口测试结果", case_rows)

    self_check_rows = [
        ["组件", "状态", "结果说明"],
        ["MySQL 数据库", "ok", "连接正常，业务主库可用"],
        ["Redis 缓存", "disabled", "本地 Redis 未启动，系统回退到内存缓存"],
        ["Tushare Pro", "error", "本地 tk.csv 令牌文件权限异常，当前环境未通过检测"],
        ["AkShare / 新浪快照", "ok", "实时快照数据源连通正常"],
        ["系统总体状态", "degraded", "部分依赖异常，但核心业务链路仍可运行"],
    ]
    add_table_after(doc, anchor, "表6-3 系统自检结果", self_check_rows)

    p63_body_1 = find_paragraph(doc, "从测试结果来看")
    set_paragraph_text(
        p63_body_1,
        "从测试结果来看，系统在股票查询、实时监控、AI 历史会话、个人研究记录管理和后台维护等方面已经形成较完整的业务闭环。页面层和接口层的关键入口均能够在当前环境下正常返回，说明系统的基础功能实现较为完整，能够满足论文所述轻量型股票量化分析系统的演示与使用需求。",
    )
    p63_body_2 = find_paragraph(doc, "同时，系统在部分外部依赖不可用时表现出一定的鲁棒性")
    set_paragraph_text(
        p63_body_2,
        "从容错表现看，当 Redis 不可用时，系统可以自动切换到内存缓存模式；当部分外部行情依赖异常时，实时监控与列表查询仍可借助其它数据源或已有缓存继续提供服务。本次系统自检虽然发现 Tushare 在当前测试环境中存在令牌文件权限问题，但该问题并未导致整站失效，反而验证了系统分层设计与多数据源回退机制的实际价值。",
    )
    p63_body_3 = find_paragraph(doc, "但也需要指出的是")
    set_paragraph_text(
        p63_body_3,
        "但也需要指出的是，当前测试仍以功能验证和接口冒烟测试为主，尚未建立完善的自动化回归测试、压力测试和端到端测试体系。后续若要进一步提升系统工程质量，还需要补充更细粒度的单元测试、接口契约测试以及针对行情高频场景的性能测试，以便在功能迭代过程中更早发现潜在问题。",
    )


def main() -> None:
    source_path = THESIS_DIR / SOURCE_NAME
    output_path = THESIS_DIR / OUTPUT_NAME

    if not source_path.exists():
        raise FileNotFoundError(f"Source thesis not found: {source_path}")

    doc = Document(str(source_path))
    update_chapter_four(doc)
    update_chapter_six(doc)
    doc.save(str(output_path))
    print(output_path)


if __name__ == "__main__":
    main()
